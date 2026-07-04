from __future__ import annotations

import argparse
import json
import math
import os
from pathlib import Path
from typing import Iterable, Sequence

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import nsdecls, qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
WORK_DIR = ROOT / ".doc_work"
ASSET_DIR = WORK_DIR / "assets"
OUTPUT_DIR = ROOT / "deliverables"
OUTPUT_PATH = OUTPUT_DIR / "VideoCompressionApp_Akademska_Tehnicka_Dokumentacija.docx"

PAGE_WIDTH_DXA = 12240
PAGE_HEIGHT_DXA = 15840
MARGIN_DXA = 1440
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGIN_DXA = {"top": 80, "bottom": 80, "start": 120, "end": 120}

NAVY = "16324F"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
TEAL = "20756E"
GOLD = "A87513"
INK = "1F2933"
MUTED = "64748B"
LIGHT_BLUE = "EAF2F8"
LIGHT_TEAL = "E8F5F2"
LIGHT_GOLD = "FBF4E5"
LIGHT_GRAY = "F4F6F9"
MID_GRAY = "D9E1E8"
WHITE = "FFFFFF"
RED = "9B1C1C"
LIGHT_RED = "FCEBEC"

FONT_REGULAR = "C:/Windows/Fonts/arial.ttf"
FONT_BOLD = "C:/Windows/Fonts/arialbd.ttf"


REFERENCES = [
    (
        "European Broadcasting Union (EBU), “EBU Core Metadata Set, Tech 3293 v1.10,” 2020.",
        "https://tech.ebu.ch/publications/tech3293",
    ),
    (
        "PBCore, “PBCore Metadata Standard: A cataloging standard and data sharing tool for audiovisual content.”",
        "https://pbcore.org/",
    ),
    (
        "Library of Congress, “PREMIS: Preservation Metadata Maintenance Activity, Version 3.0.”",
        "https://www.loc.gov/standards/premis/",
    ),
    (
        "IPTC, “NewsML-G2: Standard for exchanging text, images, video and audio news.”",
        "https://iptc.org/standards/newsml-g2/",
    ),
    (
        "MOS Project, “Media Object Server Communications Protocol.”",
        "https://mosprotocol.com/",
    ),
    (
        "Advanced Media Workflow Association (AMWA), “AS-11: Media Contribution File Formats.”",
        "https://aafassociation.org/projects/AS-11.html",
    ),
    (
        "European Broadcasting Union (EBU), “R 128: Loudness normalisation and permitted maximum level of audio signals,” Version 5.0, 2023.",
        "https://tech.ebu.ch/publications/r128",
    ),
    (
        "Library of Congress, “Recommended Formats Statement 2025–2026.”",
        "https://www.loc.gov/preservation/resources/rfs/",
    ),
    (
        "Amazon Web Services, “Media2Cloud on AWS: Ingestion workflow.”",
        "https://docs.aws.amazon.com/solutions/latest/media2cloud-on-aws/ingestion-workflow.html",
    ),
    (
        "Amazon Web Services, “Media2Cloud on AWS: Architecture details.”",
        "https://docs.aws.amazon.com/solutions/latest/media2cloud-on-aws/architecture-details.html",
    ),
    (
        "Microsoft, “Azure AI Video Indexer overview.”",
        "https://learn.microsoft.com/en-us/azure/azure-video-indexer/video-indexer-overview",
    ),
    (
        "Google Cloud, “Video Intelligence API documentation.”",
        "https://cloud.google.com/video-intelligence/docs",
    ),
    (
        "React, “React documentation.”",
        "https://react.dev/",
    ),
    (
        "MUI, “Material UI: React component library and theming system.”",
        "https://mui.com/material-ui/getting-started/",
    ),
    (
        "React Router, “React Router documentation.”",
        "https://reactrouter.com/",
    ),
    (
        "Axios, “Axios documentation.”",
        "https://axios-http.com/docs/intro",
    ),
    (
        "OpenJS Foundation, “Express 4.x API Reference.”",
        "https://expressjs.com/en/4x/api.html",
    ),
    (
        "OpenJS Foundation, “Node.js v20 Stream API.”",
        "https://nodejs.org/download/release/latest-v20.x/docs/api/stream.html",
    ),
    (
        "OpenJS Foundation, “Node.js v20 File System API: fsPromises.statfs.”",
        "https://nodejs.org/download/release/latest-v20.x/docs/api/fs.html#fspromisesstatfspath-options",
    ),
    (
        "Mongoose, “Schemas and schema indexes.”",
        "https://mongoosejs.com/docs/guide.html#indexes",
    ),
    (
        "MongoDB, “Indexes: Database Manual.”",
        "https://www.mongodb.com/docs/manual/indexes/",
    ),
    (
        "MongoDB, “The Equality, Sort, Range (ESR) Guideline.”",
        "https://www.mongodb.com/docs/manual/tutorial/equality-sort-range-guideline/",
    ),
    (
        "MongoDB, “cursor.explain() and execution statistics.”",
        "https://www.mongodb.com/docs/manual/reference/method/cursor.explain/",
    ),
    (
        "MongoDB, “TTL Indexes.”",
        "https://www.mongodb.com/docs/manual/core/index-ttl/",
    ),
    (
        "MongoDB, “Text Indexes.”",
        "https://www.mongodb.com/docs/manual/core/indexes/index-types/index-text/",
    ),
    (
        "FFmpeg Project, “FFmpeg Documentation.”",
        "https://ffmpeg.org/documentation.html",
    ),
    (
        "FFmpeg Project, “FFmpeg Formats Documentation: HLS muxer.”",
        "https://ffmpeg.org/ffmpeg-formats.html#hls-2",
    ),
    (
        "FFmpeg Project, “FFmpeg Filters Documentation: split and scale filters.”",
        "https://ffmpeg.org/ffmpeg-filters.html",
    ),
    (
        "NVIDIA, “Using FFmpeg with NVIDIA GPU Hardware Acceleration,” Video Codec SDK 13.0.",
        "https://docs.nvidia.com/video-technologies/video-codec-sdk/13.0/ffmpeg-with-nvidia-gpu/index.html",
    ),
    (
        "NVIDIA, “Video Encode and Decode GPU Support Matrix.”",
        "https://developer.nvidia.com/video-encode-and-decode-gpu-support-matrix-new",
    ),
    (
        "R. Pantos and W. May, “HTTP Live Streaming,” RFC 8216, IETF, 2017.",
        "https://datatracker.ietf.org/doc/html/rfc8216",
    ),
    (
        "Video Dev, “hls.js: HLS client based on HTML5 video and Media Source Extensions.”",
        "https://github.com/video-dev/hls.js/",
    ),
    (
        "MDN Web Docs, “HTTP range requests.”",
        "https://developer.mozilla.org/en-US/docs/Web/HTTP/Guides/Range_requests",
    ),
    (
        "MDN Web Docs, “Media Source Extensions API.”",
        "https://developer.mozilla.org/en-US/docs/Web/API/Media_Source_Extensions_API",
    ),
    (
        "M. Jones, J. Bradley and N. Sakimura, “JSON Web Token (JWT),” RFC 7519, IETF, 2015.",
        "https://datatracker.ietf.org/doc/html/rfc7519",
    ),
    (
        "NIST, “Role Based Access Control.”",
        "https://csrc.nist.gov/projects/role-based-access-control",
    ),
    (
        "OWASP, “REST Security Cheat Sheet.”",
        "https://cheatsheetseries.owasp.org/cheatsheets/REST_Security_Cheat_Sheet.html",
    ),
    (
        "OWASP, “Authorization Cheat Sheet.”",
        "https://cheatsheetseries.owasp.org/cheatsheets/Authorization_Cheat_Sheet.html",
    ),
    (
        "OWASP, “Authentication Cheat Sheet.”",
        "https://cheatsheetseries.owasp.org/cheatsheets/Authentication_Cheat_Sheet.html",
    ),
    (
        "OWASP, “Logging Cheat Sheet.”",
        "https://cheatsheetseries.owasp.org/cheatsheets/Logging_Cheat_Sheet.html",
    ),
    (
        "OptimalBits, “Bull: Redis-based queue for Node.”",
        "https://github.com/OptimalBits/bull",
    ),
    (
        "Redis, “Redis persistence.”",
        "https://redis.io/docs/latest/operate/oss_and_stack/management/persistence/",
    ),
    (
        "Express.js, “Multer: Node.js middleware for multipart/form-data.”",
        "https://github.com/expressjs/multer",
    ),
    (
        "Archiver, “Archiver API documentation.”",
        "https://www.archiverjs.com/docs/archiver/",
    ),
    (
        "MDN Web Docs, “Intersection Observer API.”",
        "https://developer.mozilla.org/en-US/docs/Web/API/Intersection_Observer_API",
    ),
    (
        "MDN Web Docs, “Page Visibility API.”",
        "https://developer.mozilla.org/en-US/docs/Web/API/Page_Visibility_API",
    ),
    (
        "Testing Library, “React Testing Library documentation.”",
        "https://testing-library.com/docs/react-testing-library/intro/",
    ),
    (
        "Web Vitals, “web-vitals library.”",
        "https://github.com/GoogleChrome/web-vitals",
    ),
    (
        "Meilisearch, “Getting started: first project.”",
        "https://www.meilisearch.com/docs/learn/getting_started/quick_start",
    ),
    (
        "OpenSearch, “OpenSearch documentation: introduction.”",
        "https://docs.opensearch.org/latest/getting-started/intro/",
    ),
]


def citations(*numbers: int) -> str:
    return " ".join(f"[{number}]" for number in numbers)


def font(name: str, size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    path = FONT_BOLD if bold else FONT_REGULAR
    try:
        return ImageFont.truetype(path, size=size)
    except OSError:
        return ImageFont.load_default()


def hex_rgb(value: str) -> tuple[int, int, int]:
    value = value.lstrip("#")
    return tuple(int(value[index:index + 2], 16) for index in (0, 2, 4))


def draw_centered_text(draw: ImageDraw.ImageDraw, box, text: str, font_obj, fill, spacing=6):
    x0, y0, x1, y1 = box
    lines = []
    for raw_line in text.split("\n"):
        words = raw_line.split()
        current = ""
        max_width = x1 - x0 - 24
        for word in words:
            candidate = f"{current} {word}".strip()
            if draw.textbbox((0, 0), candidate, font=font_obj)[2] <= max_width:
                current = candidate
            else:
                if current:
                    lines.append(current)
                current = word
        if current:
            lines.append(current)
    line_height = draw.textbbox((0, 0), "Ag", font=font_obj)[3] + spacing
    total_height = len(lines) * line_height
    y = y0 + (y1 - y0 - total_height) / 2
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font_obj)
        x = x0 + (x1 - x0 - (bbox[2] - bbox[0])) / 2
        draw.text((x, y), line, font=font_obj, fill=fill)
        y += line_height


def rounded_box(draw, box, fill, outline, radius=16, width=2):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def arrow(draw, start, end, color, width=5, head=12):
    draw.line([start, end], fill=color, width=width)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    for offset in (2.55, -2.55):
        endpoint = (
            end[0] + head * math.cos(angle + offset),
            end[1] + head * math.sin(angle + offset),
        )
        draw.line([end, endpoint], fill=color, width=width)


def make_architecture_diagram(path: Path):
    image = Image.new("RGB", (1600, 900), "white")
    draw = ImageDraw.Draw(image)
    title_font = font("Arial", 42, True)
    box_title = font("Arial", 26, True)
    box_text = font("Arial", 20)
    draw.text((60, 40), "As-built arhitektura aplikacije", font=title_font, fill=hex_rgb(NAVY))

    layers = [
        ((90, 140, 360, 760), LIGHT_GOLD, GOLD, "TV uloge",
         "Reporter\nMontažer\nProducent\nRealizator\nArhivista\nAdministrator"),
        ((430, 140, 710, 760), LIGHT_BLUE, BLUE, "Web klijent",
         "React 18\nMUI 6\nAppShell\nRole navigacija\nGlobalni upload/download\nHLS player"),
        ((780, 140, 1070, 760), LIGHT_TEAL, TEAL, "Aplikacijski sloj",
         "Express API\nAuth + RBAC\nWorkspace endpointi\nEdit jobs i SLA\nRundown / corrections\nAudit / feedback"),
        ((1140, 140, 1510, 760), LIGHT_GRAY, DARK_BLUE, "Podaci i obrada",
         "MongoDB / Mongoose\nRedis + Bull queue\nFFmpeg / FFprobe\nVideo worker\nHLS worker\nLokalni storage"),
    ]
    for box, fill, outline, heading, body in layers:
        rounded_box(draw, box, hex_rgb(fill), hex_rgb(outline), radius=18, width=3)
        x0, y0, x1, _ = box
        draw_centered_text(draw, (x0 + 12, y0 + 18, x1 - 12, y0 + 90), heading, box_title, hex_rgb(NAVY))
        draw.line((x0 + 24, y0 + 108, x1 - 24, y0 + 108), fill=hex_rgb(MID_GRAY), width=2)
        draw_centered_text(draw, (x0 + 18, y0 + 125, x1 - 18, box[3] - 24), body, box_text, hex_rgb(INK), spacing=13)

    arrow(draw, (360, 450), (430, 450), hex_rgb(BLUE))
    arrow(draw, (710, 450), (780, 450), hex_rgb(TEAL))
    arrow(draw, (1070, 450), (1140, 450), hex_rgb(DARK_BLUE))
    draw.text((70, 825), "Pregled slojeva; strelice označavaju dominantni tok zahtjeva i podataka.", font=font("Arial", 18), fill=hex_rgb(MUTED))
    image.save(path, quality=95)


def make_workflow_diagram(path: Path):
    image = Image.new("RGB", (1800, 920), "white")
    draw = ImageDraw.Draw(image)
    title_font = font("Arial", 40, True)
    box_font = font("Arial", 23, True)
    small_font = font("Arial", 18)
    draw.text((60, 35), "End-to-end TV workflow koji aplikacija koordinira", font=title_font, fill=hex_rgb(NAVY))
    labels = [
        ("1. Ingest", "Upload, event,\nvlasnik, metadata", LIGHT_GOLD, GOLD),
        ("2. Obrada", "FFprobe, master,\npreview, HLS, slike", LIGHT_BLUE, BLUE),
        ("3. Edit job", "Brief, segmenti,\nOFF, rok, komentari", LIGHT_TEAL, TEAL),
        ("4. Montaža/QC", "Final upload,\nQC i approval", LIGHT_BLUE, BLUE),
        ("5. Rundown", "Program, datum,\nredoslijed, zamjene", LIGHT_GOLD, GOLD),
        ("6. Emitovanje", "Air paket,\nchanged signal, aired", LIGHT_TEAL, TEAL),
        ("7. Arhiva", "Review, tagovi,\nkategorije, duplikati", LIGHT_GRAY, DARK_BLUE),
    ]
    x = 55
    boxes = []
    for heading, body, fill, outline in labels:
        box = (x, 220, x + 220, 550)
        boxes.append(box)
        rounded_box(draw, box, hex_rgb(fill), hex_rgb(outline), radius=18, width=3)
        draw_centered_text(draw, (x + 12, 240, x + 208, 330), heading, box_font, hex_rgb(NAVY))
        draw_centered_text(draw, (x + 12, 330, x + 208, 530), body, small_font, hex_rgb(INK), spacing=10)
        x += 250
    for left, right in zip(boxes, boxes[1:]):
        arrow(draw, (left[2], 385), (right[0], 385), hex_rgb(BLUE), width=4, head=10)

    correction = (690, 660, 1110, 845)
    rounded_box(draw, correction, hex_rgb(LIGHT_RED), hex_rgb(RED), radius=18, width=3)
    draw_centered_text(draw, (705, 670, 1095, 750), "Correction workflow", box_font, hex_rgb(RED))
    draw_centered_text(
        draw,
        (705, 740, 1095, 830),
        "Prijava greške → urgentni correction job → novi final → zamjena u rundownu → audit trag",
        small_font,
        hex_rgb(INK),
        spacing=8,
    )
    arrow(draw, (1425, 550), (1110, 705), hex_rgb(RED), width=4, head=10)
    arrow(draw, (690, 705), (925, 550), hex_rgb(RED), width=4, head=10)
    draw.text((60, 875), "Glavni tok je linearan; correction tok vraća problematični materijal u kontrolisanu montažu.", font=font("Arial", 18), fill=hex_rgb(MUTED))
    image.save(path, quality=95)


def make_media_diagram(path: Path):
    image = Image.new("RGB", (1700, 980), "white")
    draw = ImageDraw.Draw(image)
    draw.text((60, 35), "Media pipeline i pametna reprodukcija", font=font("Arial", 42, True), fill=hex_rgb(NAVY))
    source = (70, 350, 340, 610)
    rounded_box(draw, source, hex_rgb(LIGHT_GOLD), hex_rgb(GOLD), radius=18, width=3)
    draw_centered_text(draw, (85, 370, 325, 450), "Izvorni video", font("Arial", 26, True), hex_rgb(NAVY))
    draw_centered_text(draw, (85, 450, 325, 590), "MP4 / MOV / MXF\nMTS / MKV / AVI\nmaster ili sirovina", font("Arial", 20), hex_rgb(INK), spacing=12)

    process = (430, 350, 710, 610)
    rounded_box(draw, process, hex_rgb(LIGHT_TEAL), hex_rgb(TEAL), radius=18, width=3)
    draw_centered_text(draw, (445, 370, 695, 450), "FFprobe + FFmpeg", font("Arial", 26, True), hex_rgb(NAVY))
    draw_centered_text(draw, (445, 450, 695, 590), "tehnički metadata\ntranscode/finalize\nCPU ili NVENC fallback", font("Arial", 20), hex_rgb(INK), spacing=12)
    arrow(draw, (340, 480), (430, 480), hex_rgb(BLUE))

    outputs = [
        ((820, 120, 1120, 300), "Master / MP4 fallback", "Range 206\nbrowser kompatibilnost", LIGHT_GRAY, DARK_BLUE),
        ((820, 340, 1120, 520), "HLS 720p + 480p", "master.m3u8\nsegmenti od 2/4/6 s", LIGHT_BLUE, BLUE),
        ((820, 560, 1120, 740), "Thumbnail", "lazy statična\nidentifikacija", LIGHT_GOLD, GOLD),
        ((820, 780, 1120, 940), "Scrub frameovi", "6–24 JPG\nhover pregled", LIGHT_TEAL, TEAL),
    ]
    for box, heading, body, fill, outline in outputs:
        rounded_box(draw, box, hex_rgb(fill), hex_rgb(outline), radius=16, width=3)
        draw_centered_text(draw, (box[0] + 12, box[1] + 10, box[2] - 12, box[1] + 75), heading, font("Arial", 22, True), hex_rgb(NAVY))
        draw_centered_text(draw, (box[0] + 12, box[1] + 72, box[2] - 12, box[3] - 10), body, font("Arial", 18), hex_rgb(INK), spacing=8)
        arrow(draw, (710, 480), (box[0], (box[1] + box[3]) // 2), hex_rgb(BLUE), width=4, head=10)

    player = (1260, 300, 1630, 700)
    rounded_box(draw, player, hex_rgb(LIGHT_BLUE), hex_rgb(BLUE), radius=18, width=3)
    draw_centered_text(draw, (1275, 325, 1615, 410), "VideoPlayer", font("Arial", 28, True), hex_rgb(NAVY))
    draw_centered_text(
        draw,
        (1275, 420, 1615, 675),
        "kratkotrajni media ticket\n\nHLS Auto / 720p / 480p\n\nMP4 Range fallback\n\nseek bez punog Blob downloada",
        font("Arial", 20),
        hex_rgb(INK),
        spacing=12,
    )
    arrow(draw, (1120, 430), (1260, 430), hex_rgb(TEAL), width=5, head=12)
    arrow(draw, (1120, 210), (1260, 390), hex_rgb(DARK_BLUE), width=4, head=10)
    draw.text((60, 945), "HLS je sekundarni izlaz: njegov kvar ne ruši osnovnu obradu, jer MP4 Range ostaje fallback.", font=font("Arial", 18), fill=hex_rgb(MUTED))
    image.save(path, quality=95)


def make_comparison_diagram(path: Path):
    image = Image.new("RGB", (1700, 920), "white")
    draw = ImageDraw.Draw(image)
    draw.text((60, 35), "Od mrežnog foldera ka upravljanom medijskom workflowu", font=font("Arial", 40, True), fill=hex_rgb(NAVY))
    left = (70, 150, 780, 830)
    right = (920, 150, 1630, 830)
    rounded_box(draw, left, hex_rgb(LIGHT_GRAY), hex_rgb(MUTED), radius=18, width=3)
    rounded_box(draw, right, hex_rgb(LIGHT_TEAL), hex_rgb(TEAL), radius=18, width=3)
    draw_centered_text(draw, (95, 170, 755, 260), "Klasični file share", font("Arial", 30, True), hex_rgb(MUTED))
    draw_centered_text(draw, (945, 170, 1605, 260), "VideoCompressionApp", font("Arial", 30, True), hex_rgb(NAVY))
    left_items = [
        "ime fajla nosi kontekst",
        "ručno kopiranje i dogovor",
        "nejasna verzija finala",
        "preuzimanje radi pregleda",
        "status u porukama ili Excelu",
        "ograničen audit i odgovornost",
        "folder raste bez semantičke pretrage",
    ]
    right_items = [
        "asset + metadata + vlasnik",
        "role-aware workflow i job",
        "QC / approval / aired pipeline",
        "thumbnail, scrub i HLS preview",
        "rundown, SLA i notifikacije",
        "audit, correction i feedback",
        "indeksi, filteri i paginacija",
    ]
    item_font = font("Arial", 22)
    y = 300
    for item in left_items:
        draw.ellipse((115, y + 8, 131, y + 24), fill=hex_rgb(MUTED))
        draw.text((150, y), item, font=item_font, fill=hex_rgb(INK))
        y += 72
    y = 300
    for item in right_items:
        draw.ellipse((965, y + 8, 981, y + 24), fill=hex_rgb(TEAL))
        draw.text((1000, y), item, font=item_font, fill=hex_rgb(INK))
        y += 72
    arrow(draw, (800, 485), (900, 485), hex_rgb(BLUE), width=7, head=16)
    draw_centered_text(draw, (785, 365, 915, 450), "procesna\npromjena", font("Arial", 18, True), hex_rgb(BLUE))
    image.save(path, quality=95)


def make_assets():
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    make_architecture_diagram(ASSET_DIR / "architecture.png")
    make_workflow_diagram(ASSET_DIR / "workflow.png")
    make_media_diagram(ASSET_DIR / "media_pipeline.png")
    make_comparison_diagram(ASSET_DIR / "comparison.png")


def set_cell_margins(cell, margins=CELL_MARGIN_DXA):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge in ("top", "start", "bottom", "end"):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(margins[edge]))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)
    shading.set(qn("w:val"), "clear")


def set_cell_border(cell, color="D9E1E8", size=6):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "start", "bottom", "end", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        element = borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(size))
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths_dxa: Sequence[int], indent_dxa=TABLE_INDENT_DXA):
    if sum(widths_dxa) != CONTENT_WIDTH_DXA:
        raise ValueError(f"Table widths must total {CONTENT_WIDTH_DXA}, got {sum(widths_dxa)}")
    table.autofit = False
    table.allow_autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for cell, width in zip(row.cells, widths_dxa):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_hyperlink(paragraph, text: str, url: str, color=BLUE, underline=True):
    part = paragraph.part
    relationship_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color_node = OxmlElement("w:color")
    color_node.set(qn("w:val"), color)
    rpr.append(color_node)
    if underline:
        underline_node = OxmlElement("w:u")
        underline_node.set(qn("w:val"), "single")
        rpr.append(underline_node)
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), "Calibri")
    rfonts.set(qn("w:hAnsi"), "Calibri")
    rpr.append(rfonts)
    run.append(rpr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return hyperlink


def add_field(paragraph, instruction: str, display: str = ""):
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instruction
    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = display
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char_begin, instr_text, fld_char_separate, text, fld_char_end])
    return run


def configure_styles(doc: Document):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.line_spacing = 1.05

    for name in ("Caption", "Quote"):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    styles["Caption"].font.size = Pt(9)
    styles["Caption"].font.italic = True
    styles["Caption"].font.color.rgb = RGBColor.from_string(MUTED)
    styles["Caption"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    styles["Caption"].paragraph_format.space_before = Pt(4)
    styles["Caption"].paragraph_format.space_after = Pt(8)
    styles["Quote"].font.size = Pt(10.5)
    styles["Quote"].font.italic = True
    styles["Quote"].font.color.rgb = RGBColor.from_string(DARK_BLUE)
    styles["Quote"].paragraph_format.left_indent = Inches(0.35)
    styles["Quote"].paragraph_format.right_indent = Inches(0.35)
    styles["Quote"].paragraph_format.space_after = Pt(8)


def create_numbering(doc: Document):
    numbering = doc.part.numbering_part.element

    def add_abstract(abstract_id: int, num_format: str, text: str, marker_font=None):
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        level = OxmlElement("w:lvl")
        level.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        level.append(start)
        fmt = OxmlElement("w:numFmt")
        fmt.set(qn("w:val"), num_format)
        level.append(fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), text)
        level.append(lvl_text)
        suff = OxmlElement("w:suff")
        suff.set(qn("w:val"), "tab")
        level.append(suff)
        ppr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "540")
        tabs.append(tab)
        ppr.append(tabs)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "540")
        ind.set(qn("w:hanging"), "280")
        ppr.append(ind)
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "80")
        spacing.set(qn("w:line"), "290")
        spacing.set(qn("w:lineRule"), "auto")
        ppr.append(spacing)
        level.append(ppr)
        if marker_font:
            rpr = OxmlElement("w:rPr")
            fonts = OxmlElement("w:rFonts")
            fonts.set(qn("w:ascii"), marker_font)
            fonts.set(qn("w:hAnsi"), marker_font)
            rpr.append(fonts)
            level.append(rpr)
        abstract.append(level)
        numbering.append(abstract)

    add_abstract(101, "bullet", "•", "Arial")
    add_abstract(102, "decimal", "%1.")

    def add_num(num_id: int, abstract_id: int):
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abstract_ref = OxmlElement("w:abstractNumId")
        abstract_ref.set(qn("w:val"), str(abstract_id))
        num.append(abstract_ref)
        numbering.append(num)

    add_num(201, 101)
    add_num(202, 102)
    return 201, 202


def apply_num(paragraph, num_id: int):
    ppr = paragraph._p.get_or_add_pPr()
    num_pr = ppr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        ppr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_element = OxmlElement("w:numId")
    num_id_element.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_element])


def add_list_item(doc, text: str, num_id: int, bold_prefix: str | None = None):
    paragraph = doc.add_paragraph()
    apply_num(paragraph, num_id)
    if bold_prefix and text.startswith(bold_prefix):
        lead = paragraph.add_run(bold_prefix)
        set_run_font(lead, bold=True)
        rest = paragraph.add_run(text[len(bold_prefix):])
        set_run_font(rest)
    else:
        run = paragraph.add_run(text)
        set_run_font(run)
    return paragraph


def set_paragraph_keep(paragraph, keep_with_next=False, keep_together=False):
    paragraph.paragraph_format.keep_with_next = keep_with_next
    paragraph.paragraph_format.keep_together = keep_together


def add_body(doc, text: str, bold_lead: str | None = None, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    paragraph = doc.add_paragraph()
    paragraph.alignment = alignment
    if bold_lead and text.startswith(bold_lead):
        lead = paragraph.add_run(bold_lead)
        set_run_font(lead, bold=True)
        remainder = paragraph.add_run(text[len(bold_lead):])
        set_run_font(remainder)
    else:
        run = paragraph.add_run(text)
        set_run_font(run)
    return paragraph


def add_callout(doc, title: str, body: str, fill=LIGHT_BLUE, accent=BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    set_repeat_table_header(table.rows[0])
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    set_cell_border(cell, accent, 8)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(2)
    title_run = paragraph.add_run(title.upper())
    set_run_font(title_run, size=9, color=accent, bold=True)
    body_paragraph = cell.add_paragraph()
    body_paragraph.paragraph_format.space_before = Pt(0)
    body_paragraph.paragraph_format.space_after = Pt(3)
    body_paragraph.paragraph_format.line_spacing = 1.15
    body_run = body_paragraph.add_run(body)
    set_run_font(body_run, size=10.5, color=INK)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(1)
    return table


def add_table(
    doc,
    headers: Sequence[str],
    rows: Sequence[Sequence[str]],
    widths_dxa: Sequence[int],
    header_fill=LIGHT_GRAY,
    font_size=9.2,
):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths_dxa)
    set_repeat_table_header(table.rows[0])
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        shade_cell(cell, header_fill)
        set_cell_border(cell)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(2)
        paragraph.paragraph_format.space_after = Pt(2)
        run = paragraph.add_run(header)
        set_run_font(run, size=font_size, color=NAVY, bold=True)
    for row_values in rows:
        row = table.add_row()
        for index, value in enumerate(row_values):
            cell = row.cells[index]
            set_cell_border(cell)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(1)
            paragraph.paragraph_format.space_after = Pt(1)
            paragraph.paragraph_format.line_spacing = 1.08
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if index == 0 and len(headers) > 2 else WD_ALIGN_PARAGRAPH.LEFT
            run = paragraph.add_run(str(value))
            set_run_font(run, size=font_size, color=INK)
    set_table_geometry(table, widths_dxa)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(1)
    return table


def add_caption(doc, text: str):
    paragraph = doc.add_paragraph(style="Caption")
    run = paragraph.add_run(text)
    set_run_font(run, size=9, color=MUTED, italic=True)
    return paragraph


def add_figure(doc, image_path: Path, caption: str, alt_text: str, width=6.45):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run()
    inline_shape = run.add_picture(str(image_path), width=Inches(width))
    inline_shape._inline.docPr.set("descr", alt_text)
    set_paragraph_keep(paragraph, keep_with_next=True)
    add_caption(doc, caption)


def set_page_layout(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True


def add_page_furniture(doc: Document):
    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False
    table = header.add_table(rows=1, cols=2, width=Inches(6.5))
    set_table_geometry(table, [6100, 3260], indent_dxa=0)
    set_repeat_table_header(table.rows[0])
    table.cell(0, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    table.cell(0, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    left = table.cell(0, 0).paragraphs[0].add_run("VideoCompressionApp | Akademska tehnička dokumentacija")
    set_run_font(left, size=8.5, color=MUTED, bold=True)
    right = table.cell(0, 1).paragraphs[0].add_run("Revizija 1.0 | 28.06.2026.")
    set_run_font(right, size=8.5, color=MUTED)
    for cell in table.rows[0].cells:
        for p in cell.paragraphs:
            p.paragraph_format.space_after = Pt(0)
        tc_pr = cell._tc.get_or_add_tcPr()
        borders = OxmlElement("w:tcBorders")
        for edge in ("top", "start", "bottom", "end", "insideH", "insideV"):
            node = OxmlElement(f"w:{edge}")
            node.set(qn("w:val"), "nil")
            borders.append(node)
        tc_pr.append(borders)

    footer = section.footer
    footer.is_linked_to_previous = False
    ftable = footer.add_table(rows=1, cols=2, width=Inches(6.5))
    set_table_geometry(ftable, [4680, 4680], indent_dxa=0)
    set_repeat_table_header(ftable.rows[0])
    left_p = ftable.cell(0, 0).paragraphs[0]
    left_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = left_p.add_run("Nezavisni dokument — nije dio aplikacijskog repozitorija")
    set_run_font(run, size=8, color=MUTED)
    right_p = ftable.cell(0, 1).paragraphs[0]
    right_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = right_p.add_run("Stranica ")
    set_run_font(run, size=8, color=MUTED)
    field = add_field(right_p, "PAGE", "1")
    set_run_font(field, size=8, color=MUTED)
    run = right_p.add_run(" od ")
    set_run_font(run, size=8, color=MUTED)
    field = add_field(right_p, "NUMPAGES", "1")
    set_run_font(field, size=8, color=MUTED)
    for cell in ftable.rows[0].cells:
        for p in cell.paragraphs:
            p.paragraph_format.space_after = Pt(0)
        tc_pr = cell._tc.get_or_add_tcPr()
        borders = OxmlElement("w:tcBorders")
        for edge in ("top", "start", "bottom", "end", "insideH", "insideV"):
            node = OxmlElement(f"w:{edge}")
            node.set(qn("w:val"), "nil")
            borders.append(node)
        tc_pr.append(borders)


def add_cover(doc: Document):
    for _ in range(3):
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(22)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(18)
    run = kicker.add_run("AKADEMSKA TEHNIČKA I KOMPARATIVNA DOKUMENTACIJA")
    set_run_font(run, size=10.5, color=GOLD, bold=True)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)
    title.paragraph_format.keep_with_next = True
    run = title.add_run("VideoCompressionApp")
    set_run_font(run, size=31, color=NAVY, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(6)
    run = subtitle.add_run("Integrisani sistem za distribuciju, obradu i upravljanje")
    set_run_font(run, size=15, color=DARK_BLUE)
    subtitle2 = doc.add_paragraph()
    subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle2.paragraph_format.space_after = Pt(26)
    run = subtitle2.add_run("audiovizuelnim sadržajem u televizijskom okruženju")
    set_run_font(run, size=15, color=DARK_BLUE)

    scope = doc.add_paragraph()
    scope.alignment = WD_ALIGN_PARAGRAPH.CENTER
    scope.paragraph_format.left_indent = Inches(0.6)
    scope.paragraph_format.right_indent = Inches(0.6)
    scope.paragraph_format.space_after = Pt(50)
    run = scope.add_run(
        "Analiza implementirane arhitekture, operativnih workflowa, prednosti u odnosu na "
        "klasične mrežne foldere, sigurnosnih i performansnih mehanizama, ograničenja i razvojnih pravaca"
    )
    set_run_font(run, size=11.5, color=MUTED, italic=True)

    meta_rows = [
        ("Vrsta dokumenta", "Nezavisni tehnički izvještaj / akademska dokumentacija"),
        ("Predmet analize", "Stanje aplikacije i radnog prostora na dan 28.06.2026."),
        ("Revizija", "1.0"),
        ("Jezik", "Bosanski (stručni termini zadržani gdje su standardni u TV praksi)"),
    ]
    table = doc.add_table(rows=1, cols=2)
    set_repeat_table_header(table.rows[0])
    for index, header in enumerate(("Podatak", "Vrijednost")):
        cell = table.rows[0].cells[index]
        shade_cell(cell, LIGHT_GRAY)
        set_cell_border(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(header)
        set_run_font(r, size=9.5, color=NAVY, bold=True)
    for label, value in meta_rows:
        row = table.add_row()
        p = row.cells[0].paragraphs[0]
        r = p.add_run(label)
        set_run_font(r, size=9.5, color=NAVY, bold=True)
        p = row.cells[1].paragraphs[0]
        r = p.add_run(value)
        set_run_font(r, size=9.5, color=INK)
        shade_cell(row.cells[0], LIGHT_BLUE)
        for cell in row.cells:
            set_cell_border(cell)
    set_table_geometry(table, [2200, 7160])

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(40)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run("Sarajevo, juni 2026.")
    set_run_font(run, size=10.5, color=MUTED)
    doc.add_page_break()


def add_toc(doc: Document, toc_pages: dict[str, int] | None):
    doc.add_heading("Sadržaj", level=1)
    entries = [
        ("Sažetak", 1),
        ("Abstract", 1),
        ("1. Uvod, cilj i metodologija", 1),
        ("2. Problem klasične distribucije u TV kući", 1),
        ("3. As-built arhitektura aplikacije", 1),
        ("4. Funkcionalni mehanizmi po životnom ciklusu sadržaja", 1),
        ("5. Tehnička arhitektura i stack", 1),
        ("6. Media obrada, preview, streaming i download", 1),
        ("7. Pretraga, indeksiranje i performanse", 1),
        ("8. Sigurnost, ovlaštenja i sljedivost", 1),
        ("9. Prednosti u odnosu na klasično umrežavanje", 1),
        ("10. Usklađenost sa profesionalnim obrascima", 1),
        ("11. Ograničenja, rizici i granice sistema", 1),
        ("12. Operativni model uvođenja u TV kuću", 1),
        ("13. Mjerenje učinka i kriteriji prihvata", 1),
        ("14. Preporučeni razvojni pravci", 1),
        ("15. Zaključak", 1),
        ("Dodatak A. Matrica uloga i odgovornosti", 1),
        ("Dodatak B. Sažetak aplikacijskih interfejsa i entiteta", 1),
        ("Dodatak C. Pojmovnik", 1),
        ("Reference", 1),
    ]
    for title, _ in entries:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(title)
        set_run_font(run, size=10.2, color=INK)
        page = toc_pages.get(title) if toc_pages else None
        if page:
            p.paragraph_format.tab_stops.add_tab_stop(Inches(6.15))
            p.add_run("\t")
            run = p.add_run(str(page))
            set_run_font(run, size=10.2, color=MUTED)
    doc.add_page_break()


def add_front_matter(doc: Document, toc_pages):
    doc.add_heading("Sažetak", level=1)
    add_body(
        doc,
        "VideoCompressionApp je web aplikacija namijenjena centralizaciji ingest-a, obrade, "
        "distribucije, montažnih zadataka, kontrole kvaliteta, pripreme emisije i arhivskog pregleda "
        "video materijala u televizijskoj kući. Analizirano stanje koristi React/MUI klijent, "
        "Express/Node.js API, MongoDB/Mongoose modele, Redis/Bull redove obrade i FFmpeg/FFprobe "
        "media pipeline. Sistem uvodi uloge, statusne tokove, paginirane radne prostore, HLS i MP4 "
        "Range reprodukciju, sigurnosne tickete za media i download, rundown, correction workflow, "
        "audit log, notifikacije, storage nadzor i podesive media profile."
    )
    add_body(
        doc,
        "Glavni nalaz je da aplikacija ne zamjenjuje samu mrežu, storage, profesionalni NLE ili "
        "playout automatizaciju; ona zamjenjuje neformalni organizacioni sloj koji je u klasičnom "
        "okruženju često rasut između mrežnih foldera, imena fajlova, poruka, Excel tabela i usmenih "
        "dogovora. Time se povećavaju pronalazivost, sigurnost verzije, sljedivost odgovornosti, brzina "
        "pregleda i transparentnost statusa. Arhitektonski obrazac je saglasan sa praksama file-based "
        "ingest-a, proxy generisanja i metadata-driven upravljanja opisanim u EBUCore, PBCore i AWS "
        f"Media2Cloud izvorima {citations(1, 2, 9)}."
    )
    add_callout(
        doc,
        "Ključna teza",
        "Najveća vrijednost aplikacije nije samo kompresija videa. Vrijednost nastaje kada se fajl "
        "pretvori u upravljani medijski asset sa vlasnikom, metapodacima, previewem, statusom, "
        "zadacima, odobrenjem, historijom i kontrolisanim putem do etera i arhive.",
        fill=LIGHT_TEAL,
        accent=TEAL,
    )
    keywords = doc.add_paragraph()
    keywords.paragraph_format.space_before = Pt(8)
    run = keywords.add_run("Ključne riječi: ")
    set_run_font(run, bold=True, color=NAVY)
    run = keywords.add_run(
        "broadcast workflow, media asset management, ingest, FFmpeg, HLS, MongoDB, newsroom, "
        "rundown, arhiva, QC, RBAC, audit, video streaming"
    )
    set_run_font(run)

    doc.add_heading("Abstract", level=1)
    add_body(
        doc,
        "VideoCompressionApp is a web-based system designed to centralize video ingest, processing, "
        "distribution, edit assignments, quality control, rundown preparation and archive review in "
        "a television environment. The examined implementation combines a React/MUI client, an "
        "Express/Node.js API, MongoDB/Mongoose data models, Redis/Bull processing queues and an "
        "FFmpeg/FFprobe media pipeline. Its primary organizational contribution is the replacement "
        "of implicit folder-based coordination with explicit media assets, roles, statuses, deadlines, "
        "notifications, approvals, correction requests and audit records. The system should therefore "
        "be understood as a workflow and distribution layer rather than a replacement for networking, "
        "professional editing software, playout automation or long-term preservation infrastructure."
    )
    add_toc(doc, toc_pages)


def add_main_content(doc: Document, bullet_num: int, decimal_num: int):
    doc.add_heading("1. Uvod, cilj i metodologija", level=1)
    doc.add_heading("1.1. Predmet i cilj dokumenta", level=2)
    add_body(
        doc,
        "Ovaj dokument predstavlja samostalnu akademsku i tehničku analizu aplikacije "
        "VideoCompressionApp. Nije dio izvršnog koda, korisničkog uputstva niti internog changeloga. "
        "Cilj mu je opisati kako sistem funkcioniše, koje probleme rješava u televizijskoj kući, zašto "
        "je procesno napredniji od običnog dijeljenja foldera, na koje se stručne standarde i industrijske "
        "primjere naslanja, te koje granice i rizike treba uzeti u obzir prije produkcijskog uvođenja."
    )
    doc.add_heading("1.2. Metodologija", level=2)
    add_body(
        doc,
        "Analiza je izvedena kao kombinacija pregleda izvornog koda, modela podataka, API ruta, "
        "frontend radnih prostora, aplikacijske dokumentacije i primarnih tehničkih izvora. "
        "Autor dokumenta nije pretpostavio da postojanje naziva funkcije dokazuje njenu operativnu "
        "ispravnost: gdje nije izvršen stvarni broadcast ili mrežni test, nalaz je označen kao "
        "arhitektonska sposobnost ili preporučeni QA, a ne kao certificirana produkcijska garancija."
    )
    for item in [
        "Pregledano je 14 Mongoose modela, 12 route modula, 16 servisnih modula, 11 glavnih frontend stranica i 8 test fajlova.",
        "U backend route modulima evidentirano je 135 deklaracija HTTP handlera; broj predstavlja implementacijski snapshot, a ne javni API katalog.",
        "Provjereni su ključni tokovi: ingest, obrada, edit job, final/QC, rundown, correction, arhiva, download, streaming, audit i administracija storagea.",
        "Literatura je prioritetno birana iz EBU, IETF, NIST, Library of Congress, OWASP i službenih dokumentacija korištenih tehnologija.",
    ]:
        add_list_item(doc, item, bullet_num)
    add_callout(
        doc,
        "Interpretacija usklađenosti",
        "Povezanost funkcije sa EBUCore, PBCore, PREMIS, MOS, AS-11 ili EBU R 128 znači da funkcija "
        "slijedi srodan profesionalni obrazac. Ne znači da aplikacija trenutno implementira cijeli "
        "standard, izvozi njegovu punu shemu ili posjeduje formalnu certifikaciju.",
        fill=LIGHT_GOLD,
        accent=GOLD,
    )

    doc.add_heading("2. Problem klasične distribucije u TV kući", level=1)
    doc.add_heading("2.1. File share je transport, a ne workflow", level=2)
    add_body(
        doc,
        "Mrežni folder, NAS share ili SMB server uspješno rješava osnovni transport i zajednički pristup "
        "fajlovima, ali sam po sebi ne modelira urednički kontekst. Fajl ne zna ko ga je snimio, kojem "
        "eventu pripada, da li je sirovina ili final, ko ga treba montirati, da li je prošao QC, u kojoj "
        "emisiji se nalazi i da li je zamijenjen nakon što ga je realizator preuzeo. Te informacije se "
        "zato često kodiraju u naziv fajla, hijerarhiju foldera ili ljudsku komunikaciju."
    )
    add_body(
        doc,
        "Takav model proizvodi četiri vrste operativnog gubitka: vrijeme traženja, pogrešnu verziju, "
        "nevidljiv status i slab trag odgovornosti. Sa rastom broja klipova raste i cijena neformalnog "
        "dogovora. EBUCore naglašava da su metapodaci vezivno tkivo file-based produkcije, dok PBCore "
        f"i PREMIS odvajaju opis sadržaja, instancu i događaje nad digitalnim objektom {citations(1, 2, 3)}."
    )
    add_figure(
        doc,
        ASSET_DIR / "comparison.png",
        "Slika 1. Konceptualni prijelaz sa file-share distribucije na upravljani medijski workflow.",
        "Dva uporedna panela: klasični file share i VideoCompressionApp sa assetima, statusima, previewem i auditom.",
    )
    doc.add_heading("2.2. Tipični rizici folder-orijentisanog rada", level=2)
    risks = [
        ("Identitet materijala", "Duplikati, nečitljiva imena i nejasno porijeklo."),
        ("Verzija", "Fajlovi tipa FINAL_2, FINAL_NOVO ili zamjena bez evidencije."),
        ("Pronalazivost", "Pretraga je ograničena na naziv i lokaciju foldera."),
        ("Koordinacija", "Rok, brief i dopune zavise od poruka i usmenog dogovora."),
        ("QC i eter", "Odobrenje i tehnički problem nisu vezani za asset."),
        ("Arhiva", "Kasnija rekonstrukcija emitovanog sadržaja je skupa i nepouzdana."),
        ("Odgovornost", "Brisanje, zamjena i odobrenje često nemaju audit trag."),
    ]
    add_table(doc, ["Područje", "Rizik"], risks, [2400, 6960])

    doc.add_heading("3. As-built arhitektura aplikacije", level=1)
    add_body(
        doc,
        "Aplikacija je slojevita web arhitektura. React/MUI klijent pruža role-aware radne prostore; "
        "Express API provodi autentikaciju, autorizaciju i domenske akcije; MongoDB čuva assete i "
        "workflow stanje; Redis/Bull ili lokalni red koordinira asinkronu obradu; FFmpeg/FFprobe "
        "analizira i transkodira medije; lokalni storage razdvaja raw, compressed, final, preview, "
        "HLS, thumbnail, scrub, OFF i operativne fajlove. Ovakvo razdvajanje odgovornosti podsjeća na "
        "industrijske ingest orkestracije koje razdvajaju web interfejs, metadata bazu, proxy storage "
        f"i processing tok {citations(9, 10)}."
    )
    add_figure(
        doc,
        ASSET_DIR / "architecture.png",
        "Slika 2. Slojevita as-built arhitektura VideoCompressionApp.",
        "Četiri sloja: TV uloge, React web klijent, Express aplikacijski sloj te MongoDB Redis FFmpeg i storage.",
    )
    doc.add_heading("3.1. Uloge i radni prostori", level=2)
    role_rows = [
        ("Reporter", "Novi prilog, aktivni jobovi, dopune, komentari, vlastita i TV arhiva."),
        ("Editor / VideoEditor", "Job board, materijal, download paketi, montaža, final upload i correction jobovi."),
        ("Producer", "TV biblioteka, QC/approval, rundown, zamjena i correction queue."),
        ("Realizator", "Rundown pregled, detaljni preview, air paket, prijava greške i potvrda emitovanja."),
        ("Archivist", "Review queue, metadata, tagovi, kategorije, duplikati i correction status."),
        ("Admin", "Korisnici, video i job upravljanje, profili, storage, audit, feedback i održavanje."),
    ]
    add_table(doc, ["Uloga", "Primarna odgovornost"], role_rows, [2400, 6960])
    add_body(
        doc,
        "Model ima sedam tehničkih vrijednosti role, jer su Editor i VideoEditor kompatibilne produkcijske "
        "uloge. Frontend ograničava navigaciju, a backend middleware provodi provjeru dozvoljenih rola. "
        f"To odgovara osnovnoj RBAC ideji da se privilegije dodjeljuju ulozi, a korisnik ulozi {citations(36)}."
    )

    doc.add_heading("4. Funkcionalni mehanizmi po životnom ciklusu sadržaja", level=1)
    add_figure(
        doc,
        ASSET_DIR / "workflow.png",
        "Slika 3. End-to-end tok od ingest-a do arhive sa povratnom correction granom.",
        "Sedam koraka: ingest, obrada, edit job, montaža i QC, rundown, emitovanje i arhiva, uz correction povratni tok.",
    )
    doc.add_heading("4.1. Centralni ingest i evidencija asseta", level=2)
    add_body(
        doc,
        "Reporter ili ovlaštena uloga uploaduje jedan ili više fajlova uz event, datum, lokaciju i druge "
        "podatke. Backend kreira Video zapis, zadržava putanje izvornog i izvedenih fajlova, tehničke "
        "karakteristike, processing status i vlasništvo. Multipart upload koristi Multer, a podržani "
        "formati obuhvataju tipične newsroom kontejnere kao MP4, MOV, MXF, AVI, MKV, WebM i MPEG-TS. "
        f"Ovaj obrazac je srodan Media2Cloud ingest toku koji kreira identitet asseta, izvlači tehničke podatke i generiše proxy/thumbnail {citations(9, 43)}."
    )
    for item in [
        "Batch upload se prati po fajlu; neuspjeh jednog fajla ne mora poništiti cijeli batch.",
        "Processing statusi uploaded, queued, processing, completed i failed daju operativnu vidljivost.",
        "Raw recovery manifest i orphan scan omogućavaju administrativni oporavak fajla koji postoji na disku bez potpunog DB zapisa.",
        "Raw retention politika odvaja kratkoročno zadržavanje izvora od dugoročnog finala i preview izvedenica.",
    ]:
        add_list_item(doc, item, bullet_num)

    doc.add_heading("4.2. Reporter radni prostor i edit job", level=2)
    add_body(
        doc,
        "Reporter workflow je organizovan oko priče/priloga, a ne samo oko pojedinačnog fajla. Aktivni jobovi "
        "su vidljivi na početnoj površini, dok se novi prilog kreira kroz event workspace. Job može sadržati "
        "naslov, opis, script, program, content type, rok, prioritet, segmentirane izvore, marker vremena, "
        "OFF audio, komentare i change log. Reporter može naknadno dodati klip sa servera/TV arhive ili direktno "
        "sa računara bez napuštanja aktivnog joba."
    )
    add_body(
        doc,
        "Takva organizacija prati newsroom princip prema kojem su mediji, tekst i zadaci vezani za story "
        "kontekst. MOS standard opisuje razmjenu između newsroom računarskog sistema i media servera, dok "
        "NewsML-G2 formalizuje pakovanje vijesti, metapodataka i planiranja. Aplikacija ne implementira MOS ili "
        f"NewsML-G2 protokol, ali koristi isti konceptualni pomak od fajla ka uredničkom objektu {citations(4, 5)}."
    )
    add_body(
        doc,
        "EditJob razdvaja montažni status od prisutnosti na aktivnoj radnoj površini. Status opisuje tok "
        "draft/submitted/claimed/in_edit/ready_for_qc/approved, dok workspaceState može biti active, expired, "
        "closed ili cancelled. Content type može definisati SLA i grace period; scheduler pri startu i svakih "
        "pet minuta označava dospjele aktivne jobove kao expired bez brisanja materijala. Admin može promijeniti "
        "rok, kategoriju, prioritet, montažera i lifecycle stanje, dok hard delete treba ostati blokiran za job "
        "koji ima final ili rundown/air vezu."
    )
    doc.add_heading("4.3. Komentari i notifikacije", level=2)
    add_body(
        doc,
        "Komentar reportera obavještava dodijeljenog montažera; komentar montažera obavještava reportera; "
        "komentar Producenta ili Admina može obavijestiti obje operativne strane. Autor ne dobija vlastitu "
        "notifikaciju. Notification zapis čuva primaoca, autora, job, komentar, readAt i expiresAt, a TTL indeks "
        "uklanja prolaznu notifikaciju nakon definisanog roka dok komentar ostaje u trajnom job change logu. "
        f"MongoDB TTL indeks je prikladan za podatke ograničenog životnog vijeka, uz napomenu da brisanje nije trenutačno garantovano {citations(24)}."
    )
    doc.add_heading("4.4. Production Desk: materijal, montaža i final", level=2)
    add_body(
        doc,
        "Montažer dobija paginiranu listu materijala sa filterima po eventu, lokaciji, reporteru, datumu, "
        "kategoriji, processing, QC i air statusu. Thumbnail i scrub preview smanjuju potrebu za otvaranjem "
        "ili skidanjem cijelog fajla radi identifikacije. U job detaljima sirovine su prikazane kompaktno uz "
        "preview i vezu na Video Details, a edit package može uključiti odabrane segmente, OFF i manifest."
    )
    add_body(
        doc,
        "Final upload je kontrolisana tačka prelaza iz montaže u QC/approval. Video zapis razlikuje raw/edited "
        "status, processing status, qcStatus, broadcastStatus i finalApprovalStatus. Time naziv fajla prestaje "
        "biti jedini dokaz da je materijal završna verzija."
    )
    doc.add_heading("4.5. QC i odobrenje za eter", level=2)
    add_body(
        doc,
        "Aplikacija implementira ručni statusni QC: pending, passed i failed, uz bilješke, odgovornu osobu i "
        "vrijeme. Broadcast statusi vode asset od not_ready preko qc_pending/ready_for_approval do "
        "approved_for_air, aired i archived. Ovo je važan procesni temelj, ali nije isto što i automatska "
        "tehnička validacija audio loudnessa, black/freeze detekcije ili formalnog AS-11 profila. EBU R 128 "
        "preporučuje cilj -23 LUFS i true-peak/loudness deskriptore, dok AMWA AS-11 definiše ograničene formate "
        f"za predaju gotovih programa broadcasteru {citations(6, 7)}."
    )
    doc.add_heading("4.6. Producer, rundown i realizacija", level=2)
    add_body(
        doc,
        "Producent bira program i datum, pretražuje odobrenu TV biblioteku, dodaje materijal u ShowDay rundown, "
        "mijenja redoslijed i kontrolisano zamjenjuje stavku. ShowDay sadrži producers, items, activity log i "
        "download state. Realizator dobija read-only pregled sadržaja, 16:9 player, statusne i correction podatke, "
        "air package i signal da se sadržaj promijenio nakon prethodnog downloada."
    )
    add_body(
        doc,
        "Ova funkcija ne predstavlja pun NRCS ili playout automation, ali rješava važan pred-playout sloj: "
        "koji klipovi pripadaju emisiji, kojim redom, koja verzija je odobrena i da li je paket zastario. "
        f"MOS je relevantna buduća interoperabilna granica prema profesionalnim newsroom i media server sistemima {citations(5)}."
    )
    doc.add_heading("4.7. Correction workflow", level=2)
    add_body(
        doc,
        "Realizator, arhivista ili drugi ovlašteni korisnik može označiti needs_correction uz napomenu i "
        "playhead timestamp. CorrectionRequest povezuje video, rundown stavku, izvorni job, correction job, "
        "prijavitelja, montažera, status i podatke razrješenja. Ako je odgovorni montažer poznat, sistem može "
        "automatski kreirati urgentni correction job; u suprotnom Producent ili Admin usmjerava zahtjev. "
        "Otvoren zahtjev ostaje u correction queueu dok nije resolved ili dismissed, a historija čuva ko je "
        "prijavio i ko je završio ispravku."
    )
    add_callout(
        doc,
        "Operativna prednost",
        "Greška više nije samo tag bez vlasnika. Ona postaje mjerljiv zahtjev sa statusom, rokom, odgovornim "
        "montažerom, zamjenskim finalom i tragom razrješenja.",
        fill=LIGHT_RED,
        accent=RED,
    )
    doc.add_heading("4.8. Archivist Desk", level=2)
    add_body(
        doc,
        "Arhivista radi kroz review queue, listu svih dostupnih materijala i pregled grupa duplikata. Može "
        "urediti deskriptivne metapodatke, kategoriju, tagove, review status i bilješke; vidi tehničke "
        "karakteristike i correction stanje, te može inicirati ili odbaciti correction zahtjev u okviru "
        "ovlaštenja. EBUCore i PBCore daju relevantan model za buduće standardizovanje deskriptivnih i "
        "tehničkih polja, dok PREMIS naglašava objekte, događaje, agente i prava u digitalnoj prezervaciji "
        f"{citations(1, 2, 3)}."
    )
    doc.add_heading("4.9. Admin, audit i feedback", level=2)
    add_body(
        doc,
        "Admin upravlja korisnicima, videima, jobovima, broadcast programima i kategorijama, FFmpeg/media "
        "profilima, storage pragovima, preview/HLS rebuildovima, raw oporavkom, audit logovima i feedback "
        "inboxom. AuditLog čuva akciju, izvršioca, detalje i timestamp; Feedback uvodi status, prioritet, "
        "dodjelu, interne bilješke i odgovor korisniku. OWASP ističe da kvalitetan audit i logging podržavaju "
        f"detekciju, istragu i provjeru autorizacijskih problema {citations(38, 40)}."
    )
    doc.add_heading("4.10. Storage pregled i maintenance", level=2)
    add_body(
        doc,
        "Storage scanner koristi Node.js fsPromises.statfs za kapacitet fizičkog volumena i zaseban obilazak "
        "foldera za zauzeće po kategoriji: raw, compressed, final, MP4 preview, HLS, scrub, thumbnail, OFF, "
        "temp, manifesti, logovi i aplikacijski fajlovi. Scanner ne prati symlinkove, izoluje grešku kategorije, "
        "koristi single-flight zaštitu i čuva snapshot sa desetominutnim cacheom. MongoDB dbStats prikazuje se "
        f"odvojeno kako se remote/Atlas storage ne bi pogrešno pribrajao lokalnom disku {citations(19)}."
    )
    add_body(
        doc,
        "Admin podešava warning i critical procenat slobodnog prostora. Ručni refresh pokreće background scan, "
        "dok prethodni snapshot ostaje vidljiv. Preview maintenance i HLS rebuild koriste posebne redove, "
        "versioned privremeni izlaz i zamjenu tek nakon validacije; neuspjela obrada zato ne smije ukloniti "
        "prethodni ispravan preview."
    )

    doc.add_heading("5. Tehnička arhitektura i stack", level=1)
    stack_rows = [
        ("Frontend", "React 18, React Router 6, MUI 6, Axios, hls.js", "Role-aware SPA, operativni UI, API i media playback."),
        ("Backend", "Node.js 20, Express 4", "REST API, middleware, streaming i statičko serviranje builda."),
        ("Podaci", "MongoDB, Mongoose", "Asseti, workflow stanje, indeksi, TTL i audit."),
        ("Redovi", "Bull 4 + Redis ili local mode", "Odvajanje ingest, HLS i maintenance obrade."),
        ("Media", "FFmpeg/FFprobe, NVENC opcija", "Probe, transcode, HLS, MP4, thumbnail i scrub."),
        ("Transport", "HTTP, HLS, Range, ZIP", "Preview, seek i kontrolisani download paketi."),
    ]
    add_table(doc, ["Sloj", "Tehnologije", "Uloga"], stack_rows, [1500, 2900, 4960], font_size=8.8)
    add_body(
        doc,
        "React omogućava komponentni klijent, MUI standardizovane kontrole i temu, a React Router role-gated "
        f"navigaciju {citations(13, 14, 15)}. Express pruža middleware i response/stream API, dok Node streamovi "
        f"omogućavaju slanje velikih fajlova bez učitavanja kompletnog sadržaja u memoriju {citations(17, 18)}. "
        "Mongoose modeli definišu validaciju i compound indekse, uz produkcijsko kreiranje indeksa kroz "
        f"kontrolisanu skriptu umjesto automatskog autoIndex ponašanja {citations(20)}."
    )
    doc.add_heading("5.1. Asinkrona obrada i otpornost", level=2)
    add_body(
        doc,
        "Ingest obrada, HLS obrada i preview maintenance koriste odvojene redove. U produkcijskom Redis modu "
        "Bull čuva jobove izvan web procesa; local mode služi za QA i može requeueati zapise zatečene u queued "
        "ili processing stanju kada izvor postoji. Odvojena HLS konkurentnost sprečava da veliki HLS backlog "
        "blokira osnovni ingest. Redis persistence i nadzor workera ipak ostaju infrastrukturna odgovornost "
        f"implementacije {citations(41, 42)}."
    )

    doc.add_heading("6. Media obrada, preview, streaming i download", level=1)
    add_figure(
        doc,
        ASSET_DIR / "media_pipeline.png",
        "Slika 4. Media pipeline, izvedeni asseti i ticketovani player.",
        "Izvorni video prolazi kroz FFprobe i FFmpeg te daje master ili MP4 fallback, HLS, thumbnail i scrub, koje player otvara preko media ticketa.",
    )
    doc.add_heading("6.1. FFprobe i izvedeni asseti", level=2)
    add_body(
        doc,
        "FFprobe očitava kontejner, video/audio codec, rezoluciju, bitrate, framerate, trajanje i audio "
        "karakteristike. FFmpeg zatim, zavisno od moda i kompatibilnosti izvora, izrađuje master/compressed "
        "izlaz, uslovni MP4 preview, thumbnail, scrub frameove i HLS. FFmpeg dokumentacija je autoritativna "
        f"osnova za podržane formate, filtere i muxere {citations(26, 27, 28)}."
    )
    doc.add_heading("6.2. Podesivi media profili i NVENC", level=2)
    add_body(
        doc,
        "Admin može upravljati master, MP4 preview, HLS, thumbnail i scrub profilima. Preview profili ograničavaju "
        "browser izlaz na H.264, dok master može koristiti H.264/H.265 CPU ili NVIDIA enkodere. HLS podržava "
        "720p/480p, podesive bitrateove i segment trajanje. Capability probe mora potvrditi NVENC prije aktivacije, "
        "a prepoznata GPU/encoder greška može ponoviti obradu preko libx264. NVIDIA vodič dokumentuje FFmpeg GPU "
        f"akceleraciju i zavisnost od podržanog hardvera/drivera {citations(29, 30)}."
    )
    add_body(
        doc,
        "MP4 preview politika je uslovna: postojeći H.264 8-bit yuv420p MP4 sa podržanim audiom može služiti "
        "direktno kao Range fallback, dok MOV, MXF, HEVC, 10-bit ili nepodržan audio i dalje zahtijevaju "
        "browser-kompatibilnu MP4 izvedenicu. Postojeći preview se ne briše automatski. Admin prvo radi dry-run, "
        "a cleanup ponovo provjerava validan HLS, alternativni fallback, nedijeljenu putanju i storage root prije "
        "brisanja, uz audit trag."
    )
    doc.add_heading("6.3. HLS i MP4 Range fallback", level=2)
    add_body(
        doc,
        "VideoPlayer kreira kratkotrajni media ticket. Ako je HLS spreman, Safari može koristiti native HLS, "
        "dok drugi podržani browseri koriste hls.js i Media Source Extensions. Master playlist nudi 720p i "
        "480p varijantu, a korisnik bira Auto ili ručni kvalitet. RFC 8216 opisuje master/media playliste i "
        "adaptaciju bitratea, dok hls.js implementira HLS klijent iznad HTML5 video/MSE sloja "
        f"{citations(31, 32, 34)}."
    )
    add_body(
        doc,
        "Ako HLS nedostaje ili zakaže, ticketovani fallback endpoint podržava HTTP Range i odgovara sa "
        "206 Partial Content za validan raspon. Time browser može početi playback i seek bez prethodnog "
        "Axios Blob preuzimanja cijelog videa. Range zahtjevi su namijenjeni upravo media random-access i "
        f"download manager scenarijima {citations(33)}."
    )
    doc.add_heading("6.4. Thumbnail i YouTube-like scrub preview", level=2)
    add_body(
        doc,
        "Statični thumbnail se učitava lazy kada komponenta ulazi u viewport. Manifest scrub previewa navodi "
        "broj frameova, dimenzije i trajanje, a pomjeranje kursora po X osi bira odgovarajući JPG frame. "
        "Napuštanjem sličice vraća se thumbnail. Intersection Observer omogućava da klijent ne pokreće sve "
        f"zahtjeve za elemente koji nisu vidljivi {citations(45)}."
    )
    doc.add_heading("6.5. Globalni download manager", level=2)
    add_body(
        doc,
        "Korisnički downloadi koriste DownloadTicket: single video, bulk video ZIP, edit package, OFF fajl i "
        "air package. Token se u bazi čuva kao hash, ticket ima rok i status, a browser dobija kratkotrajni URL. "
        "Globalni panel prikazuje pripremu, otvaranje, streaming, completion/abort/error i omogućava nastavak rada "
        "u drugim dijelovima aplikacije. beforeunload upozorenje štiti fazu pripreme, ali nakon handoffa fizički "
        "download vodi browserov download manager."
    )
    add_body(
        doc,
        "ZIP se streama kroz Archiver, a Node streamovi primjenjuju backpressure umjesto pravljenja kompletnog "
        f"paketa u RAM-u {citations(18, 44)}. Aplikacijski status completed znači da je server završio slanje; "
        "ne dokazuje da je browser trajno upisao fajl na disk korisnika."
    )
    doc.add_heading("6.6. Background upload i zaštita od prekida", level=2)
    add_body(
        doc,
        "BackgroundUploadContext vodi red pending/uploading/done/error, prikazuje procentualni upload progress "
        "i omogućava korisniku da nastavi raditi u drugim dijelovima interfejsa. beforeunload upozorenje se "
        "aktivira dok postoje aktivni uploadi, pa slučajno zatvaranje taba ne prolazi bez potvrde. Ovaj mehanizam "
        "ne može spriječiti prekid mreže ili gašenje računara, ali čini dugotrajni prijenos vidljivim i smanjuje "
        "nenamjerne prekide."
    )

    doc.add_heading("7. Pretraga, indeksiranje i performanse", level=1)
    doc.add_heading("7.1. Workspace endpointi i paginacija", level=2)
    add_body(
        doc,
        "Velike liste se ne učitavaju u cijelosti. Video, arhiva, edit jobs, broadcast biblioteka, feedback i "
        "audit koriste workspace endpoint koncept sa page/limit, filterima, sortiranjem, summary metrikama i "
        "facetima. Frontend debounceuje unos i šalje pretragu tek nakon minimalnog broja znakova. Thumbnail, "
        "scrub i video stream nisu dio početnog JSON payload-a."
    )
    doc.add_heading("7.2. searchText, searchPrefixes i compound indeksi", level=2)
    add_body(
        doc,
        "Video, EditJob i Feedback imaju normalizovani searchText. Video dodatno ima searchPrefixes koji za "
        "riječ insert može sadržati in, ins, inse, inser i insert. Višerječni prefix upit koristi $all, a "
        "legacy dokumenti imaju kontrolisani regex fallback dok backfill ne popuni polje. Compound indeksi "
        "pokrivaju česte kombinacije statusa, vlasnika, kategorije, roka, workspaceState i datuma."
    )
    add_body(
        doc,
        "MongoDB indeks smanjuje broj dokumenata koje query mora pregledati, ali povećava storage i trošak "
        "pisanja. Zato aplikacija definira indekse u shemama, a produkcijsko kreiranje i explain provjeru vodi "
        "posebnim skriptama. ESR smjernica pomaže redoslijedu equality, sort i range polja, dok executionStats "
        f"omogućava provjeru COLLSCAN/IXSCAN ponašanja {citations(21, 22, 23)}."
    )
    doc.add_heading("7.3. Frontend performanse kao dio UX-a", level=2)
    for item in [
        "Lazy media loading sprečava inicijalno učitavanje svih thumbnailova i previewa.",
        "Debounce ograničava broj search zahtjeva dok korisnik kuca.",
        "Polling notifikacija radi samo dok je tab vidljiv, koristeći Page Visibility API obrazac.",
        "HLS segmenti i MP4 Range smanjuju time-to-first-frame u odnosu na puni Blob download.",
        "Odvojeni HLS/maintenance queue smanjuje međusobno blokiranje produkcijskih obrada.",
        "Media ticket cache smanjuje broj MongoDB čitanja i upisa po HLS segmentu.",
    ]:
        add_list_item(doc, item, bullet_num)
    add_body(
        doc,
        f"Page Visibility API omogućava reakciju na visible/hidden stanje dokumenta {citations(46)}. Ovakve "
        "optimizacije imaju direktan UX učinak: korisnik ranije vidi sadržaj i interfejs ostaje responzivan."
    )

    doc.add_heading("8. Sigurnost, ovlaštenja i sljedivost", level=1)
    doc.add_heading("8.1. Implementirane kontrole", level=2)
    security_rows = [
        ("Autentikacija", "bcrypt hash lozinke; JWT sa rokom 1 sat; obavezan JWT_SECRET."),
        ("Autorizacija", "Frontend PrivateRoute + backend authenticateToken/authorize middleware."),
        ("Uloge", "Reporter, Editor, VideoEditor, Producer, Realizator, Archivist i Admin."),
        ("Media/download", "Kratkotrajni tokeni; u bazi se čuva SHA-256 hash; TTL čišćenje."),
        ("Putanje", "Provjere da se brisanje i HLS resurs zadržavaju unutar dozvoljenog storage root-a."),
        ("Audit", "Kritične administrativne, QC, archive i workflow akcije evidentiraju izvršioca i vrijeme."),
        ("CORS", "Eksplicitni origini i same-host pristup za backend-servirani frontend."),
    ]
    add_table(doc, ["Kontrola", "As-built mehanizam"], security_rows, [2200, 7160], font_size=9)
    add_body(
        doc,
        "JWT je standardizovan URL-safe format claimova i može biti potpisan ili MAC-ovan "
        f"{citations(35)}. Uloga middlewarea odgovara RBAC modelu {citations(36)}, dok OWASP preporučuje "
        "provjeru autorizacije na svakom resursu, siguran transport i kvalitetan logging "
        f"{citations(37, 38, 40)}."
    )
    doc.add_heading("8.2. Sigurnosne granice koje ostaju", level=2)
    add_callout(
        doc,
        "Važno",
        "Ovo nije sigurnosna certifikacija. Pregled koda potvrđuje postojanje kontrola, ali ne zamjenjuje "
        "penetracijski test, reviziju konfiguracije, upravljanje tajnama, mrežnu segmentaciju i test oporavka.",
        fill=LIGHT_RED,
        accent=RED,
    )
    for item in [
        "Aplikacija ne terminira TLS; produkcija treba HTTPS preko reverse proxyja ili sigurnog VPN/tunnel sloja.",
        "Bearer token je u localStorageu, pa XSS kompromitacija može izložiti sesiju; preporučuje se CSP i razmatranje HttpOnly cookie sesije.",
        "Nisu uočeni MFA, centralna revokacija tokena, refresh token rotacija ili enterprise identity provider integracija.",
        "Audit log nije nepromjenjiv WORM zapis; Admin može upravljati audit podacima, pa za visoku forenziku treba eksterni centralni log.",
        "CORS nije zamjena za firewall, segmentaciju, VPN i ispravnu REACT_APP_API_BASE_URL konfiguraciju.",
        "Upload i FFmpeg obrada nepoznatih fajlova zahtijevaju zakrpe, ograničenja resursa i izolaciju workera.",
    ]:
        add_list_item(doc, item, bullet_num)

    doc.add_heading("9. Prednosti u odnosu na klasično umrežavanje", level=1)
    comparison_rows = [
        ("Osnovna jedinica", "Fajl i putanja", "Asset sa metadata, statusima i historijom", "Visok"),
        ("Pronalazivost", "Naziv/folder", "Paginirana pretraga, kategorije, eventi, facet filteri", "Visok"),
        ("Pregled", "Otvaranje ili download", "Thumbnail, scrub, HLS/Range player", "Visok"),
        ("Verzija", "Konvencija imena", "Final/QC/approval/rundown status", "Visok"),
        ("Koordinacija", "Poruke i usmeni dogovor", "Edit job, brief, rok, komentar i notifikacija", "Visok"),
        ("Emitovanje", "Ručno skupljanje", "Rundown, air package i changed-since-download signal", "Visok"),
        ("Ispravke", "Novi fajl i poruka", "Correction request, urgentni job i audit razrješenja", "Visok"),
        ("Arhiva", "Folder + eventualni Excel", "Review queue, metadata, tagovi, duplikati i correction", "Visok"),
        ("Odgovornost", "Ograničena", "Audit log i change/activity logovi", "Srednje-visok"),
        ("Skaliranje liste", "Folder postaje nepregledan", "Indeksi, server-side filter i paginacija", "Visok"),
        ("Održavanje", "Ručno čišćenje diska", "Storage scan, pragovi, retencija i rebuild alati", "Srednje-visok"),
        ("Pristup", "SMB/VPN mapiranje", "Browser preko kontrolisanog HTTP(S) pristupa", "Visok"),
    ]
    add_table(
        doc,
        ["Dimenzija", "Klasični share", "VideoCompressionApp", "Dobit"],
        comparison_rows,
        [1500, 2450, 4210, 1200],
        font_size=8.1,
    )
    doc.add_heading("9.1. Zašto je prednost procesna, a ne samo tehnička", level=2)
    add_body(
        doc,
        "Klasični server može imati odličan throughput, RAID i SMB permisije, ali i dalje ne rješava "
        "uredničku semantiku. Aplikacija ne pobjeđuje file share zato što 'brže kopira' svaki fajl; ona "
        "smanjuje kognitivni i koordinacijski trošak. Korisnik vidi šta treba uraditi, koji je materijal "
        "odobren i ko je odgovoran. Time se performansa mjeri i vremenom do odluke, ne samo MB/s."
    )
    doc.add_heading("9.2. Očekivani organizacioni efekti", level=2)
    for item in [
        "manje prekida montažera radi pojašnjenja materijala i rokova;",
        "manje pogrešnih finala u pripremi emisije;",
        "brži pregled velikog broja klipova bez punog downloada;",
        "kraće vrijeme pronalaska historijskog ili odobrenog sadržaja;",
        "jasnija odgovornost kod ispravke, brisanja, QC-a i zamjene;",
        "bolja osnova za mjerenje SLA, backlog-a, grešaka i kapaciteta storagea;",
        "jedinstveniji workflow za tehnički manje iskusne reportere.",
    ]:
        add_list_item(doc, item, bullet_num)

    doc.add_heading("10. Usklađenost sa profesionalnim obrascima", level=1)
    standards_rows = [
        ("File-based ingest i proxy", "Upload, asset zapis, FFprobe, master, preview, thumbnail, HLS", "AWS Media2Cloud", "Obrazac implementiran lokalno; checksum/fixity nije potpun."),
        ("AV metadata", "Naslov, event, lokacija, ljudi, codec, statusi, tagovi", "EBUCore, PBCore", "Djelimično mapiranje; nema punog standardnog exporta."),
        ("Preservation events", "Audit, processing, archive review, correction", "PREMIS", "Konceptualno srodno; rights/fixity model nije potpun."),
        ("Newsroom objekti", "Edit job, script, segmenti, rundown", "NewsML-G2, MOS", "Nema formalnog protokola/interoperabilnosti."),
        ("Air-ready delivery", "Final, QC, approval, air package", "AS-11", "Workflow srodan; izlaz nije deklarisan AS-11 compliant."),
        ("Loudness/QC", "Manualni QC status i bilješke", "EBU R 128", "Automatsko loudness mjerenje nije dokazano implementirano."),
        ("Web streaming", "HLS 720p/480p + MP4 Range", "RFC 8216, MDN", "Tehnički obrazac implementiran."),
        ("Role security", "Role-aware UI i API authorize middleware", "NIST RBAC, OWASP", "Osnovni RBAC implementiran; enterprise kontrole ostaju."),
    ]
    add_table(
        doc,
        ["Obrazac", "Aplikacija", "Referenca", "Nivo odnosa"],
        standards_rows,
        [1600, 2850, 1700, 3210],
        font_size=7.9,
    )
    doc.add_heading("10.1. Primjeri kako sličan problem rješavaju drugi", level=2)
    add_body(
        doc,
        "AWS Media2Cloud odvaja ingest, extraction, proxy/thumbnail, metadata indeks i storage slojeve; "
        "Azure AI Video Indexer i Google Video Intelligence dodaju transcript, OCR, scene, objekte i druga "
        f"vremenski vezana saznanja {citations(9, 10, 11, 12)}. VideoCompressionApp trenutno prioritet daje "
        "operativnom newsroom workflowu i lokalnoj obradi. To je racionalan prvi korak za TV kuću koja želi "
        "urediti distribuciju prije ulaganja u AI indeksiranje."
    )
    add_body(
        doc,
        "Library of Congress RFS i PREMIS podsjećaju da access proxy nije isto što i preservation master. "
        "Zato HLS i MP4 preview ne treba tretirati kao jedinu arhivsku kopiju; original/final, format politika, "
        f"fixity i backup ostaju zasebna disciplina {citations(3, 8)}."
    )

    doc.add_heading("11. Ograničenja, rizici i granice sistema", level=1)
    limitations = [
        ("NLE", "Aplikacija ne zamjenjuje Premiere Pro, Edius, DaVinci Resolve ili Avid Media Composer."),
        ("Playout", "Rundown i air package nisu puni playout automation, channel-in-a-box ili MOS gateway."),
        ("Broadcast QC", "Manualni status ne dokazuje EBU R 128, AS-11 ili station-profile tehničku usklađenost."),
        ("Arhivska prezervacija", "Nisu kompletirani checksum/fixity, rights, migration policy i geografski odvojene kopije."),
        ("Visoka dostupnost", "Single server i lokalni storage mogu biti single point of failure bez HA/backup dizajna."),
        ("Sigurnost", "Nema dokaza o MFA, SIEM-u, eksternom IdP-u, WAF-u ili formalnom pentestu."),
        ("Performanse", "Indeksi i paginacija pomažu, ali 10k/100k test, concurrency i disk I/O moraju se mjeriti na ciljnom hardveru."),
        ("Browser download", "Aplikacija prati server handoff; završetak fizičkog upisa u browseru nije potpuno observable."),
        ("Interoperabilnost", "Nema formalnog MOS, NewsML-G2, EBUCore/PBCore ili AS-11 import/export ugovora."),
    ]
    add_table(doc, ["Granica", "Posljedica"], limitations, [1900, 7460], font_size=8.8)
    doc.add_heading("11.1. Rizik pogrešne interpretacije", level=2)
    add_body(
        doc,
        "Najveći upravljački rizik je tretirati aplikaciju kao gotovu zamjenu za cijeli broadcast lanac. "
        "Ispravno pozicioniranje je: centralni distribucijski, workflow i pregledni sloj koji se može integrisati "
        "sa storageom, NLE-ovima, playoutom i arhivskom infrastrukturom. Takvo pozicioniranje čuva realnu vrijednost "
        "sistema bez obećanja koja kod i testovi ne dokazuju."
    )

    doc.add_heading("12. Operativni model uvođenja u TV kuću", level=1)
    doc.add_heading("12.1. Faze rollouta", level=2)
    rollout_steps = [
        "Inventarisati postojeće foldere, formate, korisnike, programe, kategorije i stvarne dnevne tokove.",
        "Postaviti Node.js 20, MongoDB, FFmpeg i Redis u upravljano serversko okruženje; uvesti HTTPS/VPN i firewall pravila.",
        "Definisati role i minimalne privilegije; javnu registraciju ostaviti isključenom.",
        "Pilotirati Reporter → Editor tok na jednoj redakciji i mjeriti vrijeme pronalaska/dopune/finala.",
        "Uvesti Producer/Realizator rundown tek nakon stabilnog QC i final workflowa.",
        "Migrirati staru biblioteku u batch paketima, pokrenuti search/content-type backfill i kontrolisano kreirati indekse.",
        "Graditi HLS i scrub samo u planiranim batch paketima uz storage pragove i GPU/CPU benchmark.",
        "Uvesti Archive Desk, metadata pravila, retention/backup i odgovornost za duplikate/ispravke.",
        "Provesti simulaciju incidenta: worker pad, pun disk, nedostupan MongoDB/Redis, neuspješan download i povrat iz backupa.",
    ]
    for step in rollout_steps:
        add_list_item(doc, step, decimal_num)
    doc.add_heading("12.2. Minimalna produkcijska infrastruktura", level=2)
    infrastructure_rows = [
        ("Web/API", "Node.js 20 proces iza HTTPS reverse proxyja; health check i process manager."),
        ("Workers", "Odvojeni video, HLS i maintenance worker; kontrolisana konkurentnost."),
        ("Baza", "MongoDB backup, nadzor indeksa, kapaciteta i sporih queryja."),
        ("Queue", "Redis sa persistence/backup odlukom i nadzorom backlog-a."),
        ("Storage", "Kapacitet, IOPS, backup i restore; odvojiti master/final od izvedenih previewa."),
        ("Mreža", "Gigabit ili brža LAN osnova, VPN/tunnel za udaljeni pristup, bez javnog plain HTTP porta."),
        ("Nadzor", "Disk pragovi, failed jobs, worker heartbeat, API greške i centralni logovi."),
    ]
    add_table(doc, ["Komponenta", "Minimalni zahtjev"], infrastructure_rows, [1900, 7460], font_size=8.8)

    doc.add_heading("13. Mjerenje učinka i kriteriji prihvata", level=1)
    add_body(
        doc,
        "Akademski i poslovno opravdana procjena zahtijeva baseline prije rollouta i isto mjerenje poslije. "
        "Subjektivni dojam da je sistem 'brži' nije dovoljan. Preporučuje se najmanje četverosedmični pilot sa "
        "operativnim, tehničkim i korisničkim metrikama."
    )
    kpi_rows = [
        ("Time-to-find", "Medijan od otvaranja radnog prostora do pronalaska tačnog klipa.", "≤ 60 s za poznat event/datum"),
        ("Time-to-first-frame", "Od klika Preview do prvog prikazanog framea.", "≤ 3 s na LAN-u za HLS-ready klip"),
        ("Pogrešna verzija", "Broj slučajeva korištenog zastarjelog finala.", "Trend prema nuli"),
        ("Job SLA", "Udio aktivnih jobova završenih prije roka.", "Cilj definira redakcija"),
        ("Correction MTTR", "Vrijeme od prijave greške do resolved/zamjene.", "Mjeriti po kategoriji"),
        ("Query latency", "p50/p95 workspace odgovora sa 10k+ videa.", "p95 < 1 s na ciljnom hardveru"),
        ("Playback bytes", "Da li Network panel pokazuje m3u8/segmente ili 206 Range.", "Nema punog Blob requesta"),
        ("Failed processing", "Udio failed video/HLS poslova i retry rezultat.", "< 2% nakon stabilizacije"),
        ("Storage headroom", "Slobodan kapacitet i trend rasta po asset tipu.", "Warning > 20%, critical > 10%"),
        ("Usvojivost", "Udio dnevnih tokova izvršenih bez povratka na ad-hoc foldere.", "≥ 90% nakon pilota"),
    ]
    add_table(doc, ["KPI", "Definicija", "Primjer cilja"], kpi_rows, [1900, 4960, 2500], font_size=8.3)
    add_body(
        doc,
        f"Web Vitals biblioteka može mjeriti korisničke web performanse {citations(48)}, ali TV specifični "
        "KPI-jevi moraju uključiti media startup, seek, batch ingest, download i uredničko vrijeme do odluke."
    )

    doc.add_heading("14. Preporučeni razvojni pravci", level=1)
    roadmap_rows = [
        ("P0", "Backup i restore dokaz", "Automatizovan backup baze/storagea i redovan restore test.", "Kontinuitet poslovanja"),
        ("P0", "HTTPS i sigurnosno hardening", "Reverse proxy, CSP, rate limiting, secrets, patching i pentest.", "Smanjen napadni prostor"),
        ("P1", "Automatski tehnički QC", "EBU R 128 loudness, true peak, black/freeze, audio layout, codec profil.", "Manje tehničkih grešaka"),
        ("P1", "Fixity/checksum", "SHA-256 na ingest-u i periodična provjera master/final fajlova.", "Dokaz integriteta"),
        ("P1", "Observability", "Worker health, queue depth, p95 API, FFmpeg trajanje i centralni log.", "Brža dijagnostika"),
        ("P2", "Standardni metadata export", "EBUCore/PBCore mapiranje i rights/usage polja.", "Migracija i interoperabilnost"),
        ("P2", "MOS/NRCS integracija", "Gateway prema newsroom i playout ekosistemu.", "Manje duplog unosa"),
        ("P2", "AI indeksiranje", "Transcript/OCR/scene/object kao opcioni, mjerljivi modul.", "Dublja pretraga"),
        ("P3", "Eksterni search engine", "OpenSearch/Meilisearch tek nakon mjerenja MongoDB granice.", "Skaliranje i fuzzy search"),
        ("P3", "Cold HLS politika", "Regeneracija HLS-a za rijetko gledanu arhivu.", "Kontrola storage troška"),
    ]
    add_table(doc, ["Prioritet", "Inicijativa", "Opis", "Vrijednost"], roadmap_rows, [900, 2200, 4200, 2060], font_size=8.0)
    add_body(
        doc,
        "Azure i Google pokazuju vrijednost transkripcije, OCR-a, scene i object analitike za deep search "
        f"{citations(11, 12)}. OpenSearch i Meilisearch su relevantne opcije kada MongoDB prefix/text strategija "
        f"više ne zadovoljava zahtjeve za fuzzy search, ranking i facete {citations(49, 50)}. Uvođenje treba "
        "zasnivati na mjerenju, privatnosti i trošku, a ne na samoj dostupnosti tehnologije."
    )

    doc.add_heading("15. Zaključak", level=1)
    add_body(
        doc,
        "VideoCompressionApp je evoluirao u integrisani operativni sloj između medijskog storagea i ljudskih "
        "uloga u TV kući. Njegova najjača osobina je povezivanje asseta, metapodataka, previewa, zadatka, roka, "
        "QC-a, rundowna, ispravke, downloada i arhivskog traga u jednu radnu cjelinu. U odnosu na običan mrežni "
        "share, to donosi veću pronalazivost, jasniju verziju, manje prekida u komunikaciji, brži pregled i "
        "odgovornost koja se može naknadno analizirati."
    )
    add_body(
        doc,
        "Sistem je tehnički smislen kao lokalni newsroom/MAM-lite workflow, ali nije potpuna zamjena za NLE, "
        "playout, enterprise MAM, formalni broadcast QC, preservation repository ili backup infrastrukturu. "
        "Akademski korektan zaključak zato nije da je aplikacija 'zamijenila broadcast sistem', nego da je "
        "digitalizovala i učinila mjerljivim ključni sloj distribucije i koordinacije koji u manjim i srednjim "
        "TV okruženjima često ostaje neformalan."
    )
    add_callout(
        doc,
        "Konačna ocjena",
        "Aplikacija ima jasnu praktičnu vrijednost za TV kuću pod uslovom da se uvede kao upravljani servis: "
        "sa definisanim rolama, profilima, backupom, sigurnim mrežnim pristupom, operativnim vlasnikom, "
        "mjerljivim KPI-jima i granicom prema NLE/playout/arhivskim sistemima.",
        fill=LIGHT_TEAL,
        accent=TEAL,
    )


def add_appendices(doc: Document, bullet_num: int):
    doc.add_page_break()
    doc.add_heading("Dodatak A. Matrica uloga i odgovornosti", level=1)
    matrix_rows = [
        ("Upload sirovine", "R/W", "—", "R", "—", "R", "R/W"),
        ("Aktivni edit job", "R/W", "R/W", "R", "—", "—", "R/W"),
        ("Final upload", "—", "R/W", "R/W", "—", "—", "R/W"),
        ("QC / approval", "R", "R/W*", "R/W", "R", "R", "R/W"),
        ("Rundown", "—", "—", "R/W", "R", "—", "R/W"),
        ("Air package / aired", "—", "—", "R/W", "R/W", "R", "R/W"),
        ("Correction prijava", "R", "R/W", "R/W", "R/W", "R/W", "R/W"),
        ("Archive metadata", "R", "R", "R", "R", "R/W", "R/W"),
        ("Storage / profili", "—", "—", "—", "—", "—", "R/W"),
        ("Audit / feedback triage", "R", "R", "R", "R", "R", "R/W"),
    ]
    add_table(
        doc,
        ["Funkcija", "Reporter", "Editor", "Producer", "Realizator", "Archivist", "Admin"],
        matrix_rows,
        [2180, 1196, 1196, 1196, 1196, 1196, 1200],
        font_size=7.6,
    )
    add_body(
        doc,
        "Legenda: R = pregled/čitanje; W = izmjena/akcija; — = nema primarne funkcije. Zvjezdica označava "
        "da detaljno ovlaštenje zavisi od rute i konkretnog workflow stanja. Matrica je operativni sažetak, "
        "a backend authorize pravila ostaju izvor istine."
    )

    doc.add_heading("Dodatak B. Sažetak aplikacijskih interfejsa i entiteta", level=1)
    entity_rows = [
        ("Video", "Asset, tehnički metadata, statusi, putanje, preview/HLS, QC, archive i correction."),
        ("EditJob", "Brief, script, rok/SLA, segmenti, OFF, komentari, change log i download state."),
        ("ShowDay", "Program, datum, rundown stavke, activity log, aired i download state."),
        ("CorrectionRequest", "Prijava, timestamp, izvorni/correction job, dodjela i razrješenje."),
        ("Notification", "Unread komentar signal sa TTL rokom."),
        ("DownloadTicket", "Vrsta paketa, payload, status i kratkotrajni hashovani token."),
        ("MediaTicket", "Kratkotrajni HLS/MP4 playback token vezan za video i korisnika."),
        ("Feedback", "Prijava, prioritet, status, dodjela, odgovor i komentari."),
        ("AuditLog", "Akcija, izvršilac, detalji i timestamp."),
        ("FfmpegSettings", "Master i preview profili, verzije, NVENC probe i fallback postavke."),
        ("StorageSettings", "Warning/critical pragovi za slobodan prostor."),
    ]
    add_table(doc, ["Entitet", "Odgovornost"], entity_rows, [2500, 6860], font_size=8.6)

    endpoint_rows = [
        ("Video i arhiva", "/api/videos/*, /api/archive/*", "Workspace liste, preview, scrub, stream, metadata i review."),
        ("Edit jobs", "/api/edit-jobs/*", "Kreiranje, dopuna, komentar, status, package, final i Admin lifecycle."),
        ("Broadcast", "/api/broadcast/*", "Programi, biblioteka, final approval, rundown, air i correction report."),
        ("Media", "/api/media/*", "Ticket, HLS manifest/segment i MP4 Range fallback."),
        ("Download", "/api/downloads/*", "Ticket creation, status i browser handoff."),
        ("Corrections", "/api/corrections/*", "Queue, ensure, route, claim, status i dismiss."),
        ("Admin", "/api/admin/*", "Korisnici, jobs, storage, profili, preview/HLS, audit i maintenance."),
        ("Feedback/notifikacije", "/api/feedback/*, /api/notifications/*", "Triage, komentari i unread state."),
    ]
    add_table(doc, ["Domena", "Prefiks", "Namjena"], endpoint_rows, [1800, 2700, 4860], font_size=8.3)
    add_callout(
        doc,
        "Snapshot",
        "U trenutnom radnom stanju evidentirano je 135 deklaracija route handlera. Ova brojka nije stabilan "
        "javni ugovor: kompatibilnost se određuje dokumentovanim endpointima, parametrima i response shapeom.",
        fill=LIGHT_GRAY,
        accent=DARK_BLUE,
    )

    doc.add_heading("Dodatak C. Pojmovnik", level=1)
    glossary_rows = [
        ("Asset", "Video zapis kao upravljana cjelina: fajlovi, metadata, statusi i historija."),
        ("Ingest", "Kontrolisani unos medija u sistem uz kreiranje zapisa i pokretanje obrade."),
        ("Proxy / preview", "Izvedena verzija za brz pregled; nije nužno arhivski ili air-ready master."),
        ("HLS", "HTTP Live Streaming: playlist i segmenti koji omogućavaju adaptivni playback."),
        ("Range request", "HTTP zahtjev za dio fajla, tipično uz 206 Partial Content."),
        ("Scrub preview", "Niz frameova koji se bira pomjeranjem kursora preko thumbnaila."),
        ("Edit job", "Zadatak montaže sa briefom, segmentima, rokom, komentarima i statusom."),
        ("Rundown", "Redoslijed stavki emisije za određeni program i datum."),
        ("QC", "Quality Control: urednička/tehnička provjera ispravnosti materijala."),
        ("OFF", "Audio naracija/voice-over fajl pridružen edit jobu."),
        ("SLA", "Dogovoreni rok ili pravilo vremena obrade za kategoriju joba."),
        ("Correction request", "Pratljiva prijava greške koja ostaje otvorena do razrješenja."),
        ("RBAC", "Kontrola pristupa zasnovana na ulozi."),
        ("TTL indeks", "Indeks koji omogućava automatsko uklanjanje prolaznih dokumenata nakon roka."),
        ("NVENC", "NVIDIA hardverski video encoder korišten kroz FFmpeg kada capability probe uspije."),
        ("MAM", "Media Asset Management: upravljanje medijskim assetima, metadata i workflowom."),
        ("NRCS", "Newsroom Computer System: sistem za planiranje, story i rundown tokove."),
        ("MOS", "Standard komunikacije između NRCS-a i media object servera."),
    ]
    add_table(doc, ["Pojam", "Značenje u ovom dokumentu"], glossary_rows, [2100, 7260], font_size=8.6)


def add_references(doc: Document):
    doc.add_page_break()
    doc.add_heading("Reference", level=1)
    intro = doc.add_paragraph()
    intro.paragraph_format.space_after = Pt(10)
    run = intro.add_run(
        "Svi web izvori pristupljeni su 28.06.2026. Reference su birane prvenstveno iz službenih "
        "standarda, institucija i dokumentacija proizvođača/održavatelja tehnologije."
    )
    set_run_font(run, size=10, color=MUTED, italic=True)
    for index, (label, url) in enumerate(REFERENCES, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.32)
        p.paragraph_format.first_line_indent = Inches(-0.32)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.08
        run = p.add_run(f"[{index}] {label} ")
        set_run_font(run, size=9.2, color=INK)
        add_hyperlink(p, url, url, color=BLUE, underline=True)


def set_core_properties(doc: Document):
    props = doc.core_properties
    props.title = "VideoCompressionApp — Akademska tehnička i komparativna dokumentacija"
    props.subject = "Arhitektura, TV workflow, prednosti, ograničenja i stručne reference"
    props.author = "Nezavisna tehnička dokumentacija"
    props.last_modified_by = "Nezavisna tehnička dokumentacija"
    props.keywords = "VideoCompressionApp, TV, broadcast, ingest, MAM, HLS, FFmpeg, MongoDB, rundown, QC"
    props.comments = "Nezavisni dokument; nije dio aplikacijskog izvornog koda."


def add_document_update_fields(doc: Document):
    settings = doc.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def parse_args():
    parser = argparse.ArgumentParser()
    parser.add_argument("--toc-pages", type=Path)
    parser.add_argument("--output", type=Path, default=OUTPUT_PATH)
    return parser.parse_args()


def build_document(output: Path, toc_pages: dict[str, int] | None):
    make_assets()
    output.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    set_page_layout(doc)
    configure_styles(doc)
    bullet_num, decimal_num = create_numbering(doc)
    add_page_furniture(doc)
    add_cover(doc)
    add_front_matter(doc, toc_pages)
    add_main_content(doc, bullet_num, decimal_num)
    add_appendices(doc, bullet_num)
    add_references(doc)
    set_core_properties(doc)
    add_document_update_fields(doc)
    doc.save(output)


def main():
    args = parse_args()
    toc_pages = None
    if args.toc_pages and args.toc_pages.exists():
        toc_pages = json.loads(args.toc_pages.read_text(encoding="utf-8"))
    build_document(args.output, toc_pages)
    print(args.output)


if __name__ == "__main__":
    main()
