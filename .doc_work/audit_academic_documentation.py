from __future__ import annotations

import re
import sys
import zipfile
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "deliverables" / "VideoCompressionApp_Akademska_Tehnicka_Dokumentacija.docx"
EXPECTED_WIDTH = 9360
EXPECTED_REFERENCES = 50


def fail(message: str):
    print(f"FAIL: {message}")
    sys.exit(1)


def int_attr(node, name: str) -> int:
    value = node.get(qn(name))
    return int(value) if value is not None else -1


def audit_tables(doc: Document):
    for table_index, table in enumerate(doc.tables, start=1):
        tbl_pr = table._tbl.tblPr
        tbl_w = tbl_pr.first_child_found_in("w:tblW")
        if tbl_w is None or int_attr(tbl_w, "w:w") != EXPECTED_WIDTH:
            fail(f"table {table_index} tblW is not {EXPECTED_WIDTH}")
        grid = table._tbl.tblGrid
        widths = [int_attr(col, "w:w") for col in grid]
        if sum(widths) != EXPECTED_WIDTH:
            fail(f"table {table_index} grid width sum is {sum(widths)}")
        for row_index, row in enumerate(table.rows, start=1):
            if len(row.cells) != len(widths):
                fail(f"table {table_index} row {row_index} cell count mismatch")
            for cell_index, (cell, expected) in enumerate(zip(row.cells, widths), start=1):
                tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
                if tc_w is None or int_attr(tc_w, "w:w") != expected:
                    fail(
                        f"table {table_index} row {row_index} cell {cell_index} "
                        f"width mismatch"
                    )
        for row in table.rows:
            tr_height = row._tr.get_or_add_trPr().find(qn("w:trHeight"))
            if tr_height is not None and tr_height.get(qn("w:hRule")) == "exact":
                fail(f"table {table_index} contains fixed exact row height")


def audit_package():
    with zipfile.ZipFile(DOCX) as package:
        document_xml = package.read("word/document.xml").decode("utf-8")
        styles_xml = package.read("word/styles.xml").decode("utf-8")
        numbering_xml = package.read("word/numbering.xml").decode("utf-8")
        rels_xml = package.read("word/_rels/document.xml.rels").decode("utf-8")
        core_xml = package.read("docProps/core.xml").decode("utf-8")

    if "�" in document_xml:
        fail("replacement character found in document XML")
    placeholder_patterns = (
        r"\bTODO\b",
        r"\bTBD\b",
        r"\[INSERT",
        r"Lorem ipsum",
    )
    for placeholder in placeholder_patterns:
        if re.search(placeholder, document_xml, flags=re.IGNORECASE):
            fail(f"placeholder found: {placeholder}")

    if document_xml.count("<w:drawing>") != 4:
        fail("expected exactly four inline figures")
    if document_xml.count('descr="') < 4:
        fail("one or more figures are missing alt descriptions")
    if rels_xml.count('TargetMode="External"') != EXPECTED_REFERENCES:
        fail("external hyperlink count does not match reference count")
    if 'w:abstractNumId="101"' not in numbering_xml or 'w:abstractNumId="102"' not in numbering_xml:
        fail("custom bullet/decimal numbering definitions are missing")
    if 'dc:title' not in core_xml or "Akademska tehnička" not in core_xml:
        fail("core document title metadata is missing")
    if 'w:pgSz w:w="12240" w:h="15840"' not in document_xml:
        fail("Letter page size is missing")
    if 'w:pgMar w:top="1440"' not in document_xml:
        fail("1-inch page margins are missing")
    if 'w:styleId="Heading1"' not in styles_xml or 'w:styleId="Heading2"' not in styles_xml:
        fail("heading styles are missing")


def main():
    if not DOCX.exists():
        fail(f"missing deliverable: {DOCX}")
    if ROOT / "docs" in DOCX.parents:
        fail("deliverable must remain outside the application docs folder")

    doc = Document(DOCX)
    audit_tables(doc)
    audit_package()

    paragraphs = [paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip()]
    headings_1 = [p.text for p in doc.paragraphs if p.style.name == "Heading 1"]
    headings_2 = [p.text for p in doc.paragraphs if p.style.name == "Heading 2"]
    if len(headings_1) != 22 or len(headings_2) != 33:
        fail(f"unexpected heading counts: H1={len(headings_1)} H2={len(headings_2)}")
    required = [
        "Sažetak",
        "Abstract",
        "9. Prednosti u odnosu na klasično umrežavanje",
        "11. Ograničenja, rizici i granice sistema",
        "Dodatak A. Matrica uloga i odgovornosti",
        "Reference",
    ]
    for heading in required:
        if heading not in headings_1:
            fail(f"required heading missing: {heading}")

    full_text = "\n".join(paragraphs)
    reference_labels = {
        int(match.group(1))
        for match in re.finditer(r"^\[(\d+)\]\s", full_text, flags=re.MULTILINE)
    }
    if reference_labels != set(range(1, EXPECTED_REFERENCES + 1)):
        fail("reference labels are incomplete or non-sequential")
    citations = [int(value) for value in re.findall(r"\[(\d+)\]", full_text)]
    if not citations or max(citations) > EXPECTED_REFERENCES:
        fail("invalid in-text citation detected")

    file_size = DOCX.stat().st_size
    if file_size < 250_000:
        fail(f"document file unexpectedly small: {file_size}")

    print(
        "PASS "
        f"size={file_size} "
        f"paragraphs={len(paragraphs)} "
        f"H1={len(headings_1)} "
        f"H2={len(headings_2)} "
        f"tables={len(doc.tables)} "
        f"references={len(reference_labels)} "
        f"citations={len(citations)}"
    )


if __name__ == "__main__":
    main()
